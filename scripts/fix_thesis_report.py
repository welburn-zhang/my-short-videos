import re
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path("/Users/zhangwenbo/Desktop/邵艺博毕业论文终稿.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_修改前备份.docx")


def backup_file():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)


def iter_paragraphs(parent):
    for p in parent.paragraphs:
        yield p
    for t in parent.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in iter_paragraphs(cell):
                    yield p


def iter_all_story_paragraphs(doc):
    for p in iter_paragraphs(doc):
        yield p
    for section in doc.sections:
        for p in iter_paragraphs(section.header):
            yield p
        for p in iter_paragraphs(section.footer):
            yield p


def set_run_fonts(run):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def ensure_ppr(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    return ppr


def set_left_chars(paragraph, chars):
    ppr = ensure_ppr(paragraph)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:leftChars"), str(int(chars * 100)))
    ind.attrib.pop(qn("w:left"), None)


def clear_left_chars(paragraph):
    ppr = ensure_ppr(paragraph)
    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        ind.attrib.pop(qn("w:leftChars"), None)
        ind.attrib.pop(qn("w:firstLineChars"), None)
        ind.attrib.pop(qn("w:hangingChars"), None)


def set_hanging_chars(paragraph, chars):
    ppr = ensure_ppr(paragraph)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    val = str(int(chars * 100))
    ind.set(qn("w:leftChars"), val)
    ind.set(qn("w:hangingChars"), val)
    ind.attrib.pop(qn("w:left"), None)
    ind.attrib.pop(qn("w:hanging"), None)


def strip_leading_spaces(paragraph):
    for run in paragraph.runs:
        if run.text:
            new_text = run.text.lstrip()
            if new_text != run.text:
                run.text = new_text
            break


def set_paragraph_text(paragraph, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    first = runs[0]
    first.text = new_text
    for run in runs[1:]:
        run.text = ""


def dominant_run_size(paragraph):
    sized = []
    for run in paragraph.runs:
        if run.font.size and run.text.strip():
            sized.append((run.font.size, len(run.text.strip())))
    if not sized:
        return None
    counter = Counter()
    for size, weight in sized:
        counter[size] += weight
    return counter.most_common(1)[0][0]


def normalize_mixed_run_sizes(paragraph):
    dominant = dominant_run_size(paragraph)
    if dominant is None:
        return
    sizes = {run.font.size for run in paragraph.runs if run.font.size}
    if len(sizes) <= 1:
        return
    total = sum(len(run.text.strip()) for run in paragraph.runs if run.text.strip())
    dom = sum(len(run.text.strip()) for run in paragraph.runs if run.font.size == dominant and run.text.strip())
    if not total or dom / total < 0.65:
        return
    for run in paragraph.runs:
        if run.text.strip():
            run.font.size = dominant


def move_after(paragraph_to_move, target_paragraph):
    target_paragraph._p.addnext(paragraph_to_move._p)


def fix_toc_paragraph(paragraph):
    style = paragraph.style.name.lower()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if style == "toc 1":
        set_left_chars(paragraph, 0)
    elif style == "toc 2":
        set_left_chars(paragraph, 1)
    elif style == "toc 3":
        set_left_chars(paragraph, 2)


def fix_heading_paragraph(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    clear_left_chars(paragraph)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_with_next = True


def first_chinese_period_to_dot(text):
    return text.replace("。", ".", 1) if "。" in text else text


def fix_reference_entry(text):
    text = first_chinese_period_to_dot(text)
    replacements = {
        "暨南大学,2018.": "广州:暨南大学,2018.",
        "山西财经大学,2018.": "太原:山西财经大学,2018.",
        "江西师范大学,2025.": "南昌:江西师范大学,2025.",
        "南京大学,2018.": "南京:南京大学,2018.",
        "广西师范大学,2024.": "桂林:广西师范大学,2024.",
        "西北工业大学,2017.": "西安:西北工业大学,2017.",
        "杭州电子科技大学,2016.": "杭州:杭州电子科技大学,2016.",
        "广东外语外贸大学,2025.": "广州:广东外语外贸大学,2025.",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    if text.startswith("加快提升全民数字素养与技能为网络强国建设提供有力支撑"):
        text = (
            "中央网信办信息化发展局.加快提升全民数字素养与技能为网络强国建设提供有力支撑"
            "[J].中国网信,2025,(08):18-21."
        )
    return text


def main():
    backup_file()
    doc = Document(str(DOC_PATH))

    body_paragraphs = doc.paragraphs

    # Global font normalization for mixed Chinese/English content.
    for p in iter_all_story_paragraphs(doc):
        for run in p.runs:
            set_run_fonts(run)

    # Fix TOC formatting in-place.
    for p in body_paragraphs:
        if p.style.name.lower().startswith("toc "):
            fix_toc_paragraph(p)

    # Update two TOC page numbers called out by the report.
    if len(body_paragraphs) > 108:
        for run in body_paragraphs[103].runs:
            if run.text == "27":
                run.text = "26"
        for run in body_paragraphs[108].runs:
            if run.text == "37":
                run.text = "36"

    # English abstract block.
    for idx in [49, 50, 51, 52, 54]:
        p = body_paragraphs[idx]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        strip_leading_spaces(p)
        for run in p.runs:
            if run.text:
                run.text = re.sub(r" {2,}", " ", run.text)
                run.font.size = Pt(12)

    # Remove orphan headings / normalize numbered headings and key section titles.
    heading_re = re.compile(r"^\d+(?:\.\d+){0,3}\s")
    special_titles = {"摘要", "Abstract", "目录", "参考文献", "致谢", "附录：调查问卷"}
    for p in body_paragraphs:
        txt = p.text.strip()
        if txt in special_titles or (heading_re.match(txt) and len(txt) <= 40):
            fix_heading_paragraph(p)

    # Specific numbered headings with extra spaces / punctuation.
    body_paragraphs[301].text = "4.2.1 AI 求职使用反馈量表"
    body_paragraphs[305].text = "4.2.2 数字素养量表"
    body_paragraphs[463].text = "5.1.5 “是否使用AI”与“如何使用AI”：从二元区分走向使用质量"
    for idx in [301, 305, 463]:
        fix_heading_paragraph(body_paragraphs[idx])
        for run in body_paragraphs[idx].runs:
            set_run_fonts(run)

    # Replace half-width quotes in Chinese paragraph.
    body_paragraphs[400].text = body_paragraphs[400].text.replace('"双一流"', "“双一流”")
    for run in body_paragraphs[400].runs:
        set_run_fonts(run)

    # Keep figure/table titles with the following object and normalize spacing.
    caption_prefixes = ("图", "Figure", "表", "Table")
    for p in body_paragraphs:
        txt = p.text.strip()
        if txt.startswith(caption_prefixes):
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            strip_leading_spaces(p)

    # Move Figure 4.1 titles below the image.
    move_after(body_paragraphs[380], body_paragraphs[382])
    move_after(body_paragraphs[381], body_paragraphs[380])
    if body_paragraphs[381].text.endswith("."):
        body_paragraphs[381].text = body_paragraphs[381].text[:-1]
        for run in body_paragraphs[381].runs:
            set_run_fonts(run)

    # Paragraphs explicitly mentioned for spacing/format consistency.
    for idx in [209, 210, 298, 299, 359, 360, 380, 381, 432, 433]:
        p = body_paragraphs[idx]
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True

    # Remove leading spaces before English table titles.
    for idx in [390, 433]:
        strip_leading_spaces(body_paragraphs[idx])

    # Reference heading should start on a new page.
    prev = body_paragraphs[509]
    if prev.runs:
        prev.runs[-1].add_break(WD_BREAK.PAGE)
    else:
        prev.add_run().add_break(WD_BREAK.PAGE)
    body_paragraphs[510].paragraph_format.keep_with_next = True

    # Reference entries: format and normalize GB/T 7714 details.
    for idx in range(511, 554):
        p = body_paragraphs[idx]
        new_text = fix_reference_entry(p.text.strip())
        set_paragraph_text(p, new_text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        set_hanging_chars(p, 2.02)
        for run in p.runs:
            run.font.size = Pt(10.5)
            set_run_fonts(run)

    # Normalize paragraphs where a few characters were accidentally smaller.
    for p in body_paragraphs:
        normalize_mixed_run_sizes(p)

    # Final pass for leading spaces and quote/title cleanup.
    for p in body_paragraphs:
        strip_leading_spaces(p)

    doc.save(str(DOC_PATH))
    print(f"Saved: {DOC_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
