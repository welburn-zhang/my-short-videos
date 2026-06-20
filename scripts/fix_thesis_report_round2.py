from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path("/Users/zhangwenbo/Desktop/邵艺博毕业论文终稿.docx")


ASCII_EXTRA = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789%+-=<>./,:;!?()[]{}'\"@&_#*~`|\\ αβγΔδμστω–—")


def remove_all_runs(paragraph):
    for r in list(paragraph._p.findall(qn("w:r"))):
        paragraph._p.remove(r)


def set_run_font(run, east_asia="宋体", size=Pt(12), bold=False):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.name = "Times New Roman"
    run.font.size = size
    run.bold = bold


def classify_char(ch):
    if ch == "\t":
        return "tab"
    if ch == "\n":
        return "newline"
    if ch == " ":
        return "space"
    if ch in ASCII_EXTRA:
        return "ascii"
    code = ord(ch)
    if 0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF:
        return "cjk"
    if 0xFF00 <= code <= 0xFFEF:
        return "cjk"
    if ch in "，。；：！？、“”‘’（）《》【】—…·":
        return "cjk"
    return "ascii"


def rebuild_paragraph(paragraph, east_asia="宋体", size=Pt(12), bold=False):
    text = paragraph.text
    remove_all_runs(paragraph)
    if not text:
        return
    current_kind = None
    buf = []

    def flush():
        nonlocal buf, current_kind
        if not buf:
            return
        run = paragraph.add_run("".join(buf))
        if current_kind != "tab":
            set_run_font(run, east_asia=east_asia if current_kind != "ascii" else east_asia, size=size, bold=bold)
        buf = []

    for ch in text:
        kind = classify_char(ch)
        if kind == "tab":
            flush()
            run = paragraph.add_run("\t")
            set_run_font(run, east_asia=east_asia, size=size, bold=bold)
            current_kind = None
            continue
        normalized = "ascii" if kind in {"ascii", "space"} else "cjk"
        if current_kind is None:
            current_kind = normalized
            buf.append(ch)
        elif normalized == current_kind:
            buf.append(ch)
        else:
            flush()
            current_kind = normalized
            buf.append(ch)
    flush()


def ensure_spacing_xml(paragraph, before=0, after=0, line=300, line_rule="auto"):
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), line_rule)
    for attr in ["beforeLines", "afterLines", "beforeAutospacing", "afterAutospacing"]:
        spacing.attrib.pop(qn(f"w:{attr}"), None)


def ensure_indent_chars(paragraph, left_chars=None, first_line_chars=None, clear_left=False):
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if clear_left:
        ind.attrib.pop(qn("w:left"), None)
        ind.attrib.pop(qn("w:leftChars"), None)
    if left_chars is not None:
        ind.set(qn("w:leftChars"), str(int(left_chars * 100)))
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(int(first_line_chars * 100)))
        ind.set(qn("w:firstLine"), "480" if first_line_chars == 2 else "0")


def set_style_indent_chars(style, left_chars):
    ppr = style._element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:leftChars"), str(int(left_chars * 100)))
    ind.attrib.pop(qn("w:left"), None)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def fix_toc_entry(paragraph, left_chars):
    ensure_indent_chars(paragraph, left_chars=left_chars, clear_left=True)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = None
    ensure_spacing_xml(paragraph)


def set_single_run(paragraph, text, east_asia="宋体", size=Pt(12), bold=False):
    remove_all_runs(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)


def update_toc_page(paragraph, new_page):
    text = paragraph.text
    if "\t" not in text:
        return
    title, _ = text.rsplit("\t", 1)
    new_text = f"{title}\t{new_page}"
    set_single_run(paragraph, new_text, east_asia="宋体", size=None if False else Pt(12), bold=False)


def main():
    doc = Document(str(DOC_PATH))
    p = doc.paragraphs

    # 1) 目录缩进改成检测要求
    set_style_indent_chars(doc.styles["toc 2"], 1)
    set_style_indent_chars(doc.styles["toc 3"], 2)
    for idx in [58, 59, 62, 67, 71, 75, 80, 81, 85, 91, 99, 101, 102, 103, 104, 105, 106, 108, 109, 110]:
        if idx < len(p):
            fix_toc_entry(p[idx], 1)
    for idx in [60, 61, 63, 64, 68, 69, 70, 72, 73, 76, 77, 82, 83, 84, 86, 87, 88, 89, 90, 92, 93, 94, 95, 96, 100]:
        if idx < len(p):
            fix_toc_entry(p[idx], 2)

    # 目录页码按最新检测结果回填
    page_updates = {
        59: "2",
        103: "27",
        107: "37",
        108: "37",
        109: "40",
        111: "44",
        112: "47",
        113: "48",
    }
    for idx, page_no in page_updates.items():
        if idx < len(p):
            update_toc_page(p[idx], page_no)
            fix_toc_entry(p[idx], 1 if idx in {59, 103, 107, 108, 109, 111, 112, 113} else 2)

    # 2) 中文摘要改回楷体，英文摘要清掉段前段后残留
    for idx in [41, 42, 43, 44]:
        rebuild_paragraph(p[idx], east_asia="楷体_GB2312", size=Pt(12), bold=False)
        ensure_spacing_xml(p[idx])
        ensure_indent_chars(p[idx], first_line_chars=2)
    set_single_run(p[46], p[46].text, east_asia="黑体", size=Pt(12), bold=False)
    ensure_spacing_xml(p[46])

    for idx in [49, 50, 51, 52, 54]:
        rebuild_paragraph(p[idx], east_asia="宋体", size=Pt(12), bold=False)
        ensure_spacing_xml(p[idx])
        if idx != 54:
            ensure_indent_chars(p[idx], first_line_chars=2)

    # 3) 剩余英数字体报错段落，重建为中英分字体
    body_fix = [
        239, 296, 321, 326, 327, 330, 332, 333, 336, 339,
        400, 405, 410, 416, 417, 418, 469, 480, 490
    ]
    for idx in body_fix:
        rebuild_paragraph(p[idx], east_asia="宋体", size=Pt(12), bold=False)
        ensure_spacing_xml(p[idx])

    # 4) 标题/图题/小标题剩余字号与加粗问题
    for idx in [301, 305, 463]:
        set_single_run(p[idx], p[idx].text, east_asia="宋体", size=Pt(12), bold=True)
        ensure_spacing_xml(p[idx])
        set_keep_with_next(p[idx], True)
    set_single_run(p[382], p[382].text, east_asia="宋体", size=Pt(9), bold=True)
    ensure_spacing_xml(p[382])

    # 5) 41-42页那三段被识别成单倍行距，显式写回 1.25 倍
    for idx in [416, 417, 418]:
        ensure_spacing_xml(p[idx])

    # 6) 可能被拆坏的标题字号
    for idx in [400, 404, 409]:
        rebuild_paragraph(p[idx], east_asia="宋体", size=Pt(12), bold=False)
        ensure_spacing_xml(p[idx])

    # 7) 目录与正文连接处保留一空行
    if p[55].text != "":
        set_single_run(p[55], "", east_asia="宋体", size=Pt(12), bold=False)

    doc.save(str(DOC_PATH))
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    main()
