from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path("/Users/zhangwenbo/Desktop/邵艺博毕业论文终稿.docx")
HEADER_TEXT = "北京林业大学本科毕业论文"


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_run_fonts(run, east_asia="宋体", size=Pt(9), bold=False):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.name = "Times New Roman"
    run.font.size = size
    run.bold = bold


def set_para_spacing_single(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")


def set_bottom_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        ppr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")


def remove_borders(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is not None:
        ppr.remove(p_bdr)


def set_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    if not header.paragraphs:
        p = header.add_paragraph()
    else:
        p = header.paragraphs[0]
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_fonts(run, east_asia="宋体", size=Pt(9), bold=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing_single(p)
    set_bottom_border(p)
    for extra in header.paragraphs[1:]:
        clear_paragraph(extra)


def clear_header(section):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        clear_paragraph(p)
        remove_borders(p)


def add_page_field(paragraph):
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)
    set_run_fonts(run_begin, east_asia="宋体", size=Pt(10.5), bold=False)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._r.append(instr)
    set_run_fonts(run_instr, east_asia="宋体", size=Pt(10.5), bold=False)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    set_run_fonts(run_sep, east_asia="宋体", size=Pt(10.5), bold=False)

    run_text = paragraph.add_run("1")
    set_run_fonts(run_text, east_asia="宋体", size=Pt(10.5), bold=False)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)
    set_run_fonts(run_end, east_asia="宋体", size=Pt(10.5), bold=False)


def set_footer_with_page(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing_single(p)
    add_page_field(p)
    for extra in footer.paragraphs[1:]:
        clear_paragraph(extra)


def clear_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        clear_paragraph(p)


def set_page_number_format(section, fmt=None, start=None):
    sectPr = section._sectPr
    pg_num_type = sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sectPr.append(pg_num_type)
    if fmt is not None:
        pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))


def main():
    doc = Document(str(DOC_PATH))
    sections = doc.sections

    # Global distances per spec: 1.5 cm
    for sec in sections:
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.5)

    # Section 0: before Chinese abstract, no header/footer/page numbers
    clear_header(sections[0])
    clear_footer(sections[0])
    set_page_number_format(sections[0], fmt="decimal", start=1)

    # Section 1: Chinese abstract through table of contents
    set_header(sections[1], HEADER_TEXT)
    set_footer_with_page(sections[1])
    set_page_number_format(sections[1], fmt="upperRoman", start=1)

    # Section 2: main text onward
    set_header(sections[2], HEADER_TEXT)
    set_footer_with_page(sections[2])
    set_page_number_format(sections[2], fmt="decimal", start=1)

    doc.save(str(DOC_PATH))
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    main()
